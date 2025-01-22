from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Inches, Pt

FOLDER_PATH = 'images/'
# Укажите путь к Tesseract (только для Windows, замените путь на ваш)
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
result = ''
doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(14)

# Откройте изображение
for i in range(0, 382):
    image_path = FOLDER_PATH + str(i) + '.png'  # Замените на путь к вашему изображению
    image = Image.open(image_path)

    # Распознайте текст (указание русского языка)
    text = pytesseract.image_to_string(image, lang='rus')
    doc.add_paragraph(text)
    # doc.add_picture(FOLDER_PATH + str(i) + '.png', width=Inches(4))
    doc.add_paragraph('\n')

doc.save('result.docx')
print('Готово.')
